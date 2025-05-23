﻿# -*- coding: utf-8 -*-
import sys # даёт доступ к данным из командной строки которую мы запускаем в c++
from docx import Document # доступ к файлам типа .docx
from docx.shared import Pt # доступ к размеру шрифта
from docx.shared import Mm # доступ к миллиметрам
from docx.enum.text import WD_LINE_SPACING # доступ к межстрочному интервалу
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # доступ к выравниванию текста
from docx2pdf import convert # доступ к конвертации в pdf
import win32com.client # сильная библиотека по word
import argparse # для удобной работы с данными из с++

def DocxFormat(input_path, output_path, extra_fields=None, use_pdf=False, use_check=False, app_data=None):
    doc = Document(input_path) # создаём переменную на основе docx документа и его расположения

    if not use_pdf and not use_check:
        sections = doc.sections # .sections хранит информацию о полях
        for section in sections:
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(30)
            section.right_margin = Mm(10)
        # изменяем поля под нужный формат

        for para in doc.paragraphs: # проходим по абзацам
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # Междустрочный интервал
            para.paragraph_format.first_line_indent = Mm(12.5) # Абзацный отступ
            para.paragraph_format.space_before = Pt(0)  # Убирать отступ сверху
            para.paragraph_format.space_after = Pt(0)   # Убирать отступ снизу
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # Выравнивание по ширене

            for run in para.runs: # разбиваем абзацы на блоки с одинаковым форматирование
                run.font.name = 'Times New Roman' 
                run.font.size = Pt(14)
            # font хранит в себе настройки шрифта и через разные методы мы их изменяем

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)

    titleParagraphsCount = 0 # Счётчик для количества строк в титульном листе

    if extra_fields:
        title = Document()

        def add_center(text="", font_size=12):
            p = title.add_paragraph(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)

        def add_right(text="", font_size=14):
            p = title.add_paragraph(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)

        add_center("Министерство науки и высшего образования Российской Федерации")
        add_center("Федеральное государственное бюджетное ")
        add_center("образовательное учреждение высшего образования")
        add_center("«Новгородский государственный университет имени Ярослава Мудрого»")
        add_center(extra_fields[0])
        add_center(" ", 12)
        add_center(extra_fields[1])
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(extra_fields[2], 16)
        add_center(" ", 14)
        add_center(" ", 14)
        add_center(extra_fields[3] + " по учебной дисциплине «" + extra_fields[4] + 
                   "» по специальности " + extra_fields[5] + " " + extra_fields[6], 14)
        add_center(extra_fields[10] + "." + extra_fields[11] + " " + extra_fields[8] + " " + extra_fields[12] + extra_fields[13], 14)

        title.add_paragraph()
        title.add_paragraph()

        add_right("Руководитель     ")
        add_right("______________/ " + extra_fields[9])
        add_right("«___»_________________2025 г.")
        add_right(" ", 14)
        add_right(" ", 14)
        add_right("Студент группы " + extra_fields[8])
        add_right("______________/ " + extra_fields[7])
        add_right("«___»_________________2025 г.")

        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()
        title.add_paragraph()

        titleParagraphsCount = len(title.paragraphs)

        for element in doc.element.body[:]:
            title.element.body.append(element)
        doc._body.clear_content()
        for element in title.element.body[:]:
            doc.element.body.append(element)

    if not use_pdf and not use_check:
        doc.save(output_path)  # Сохраняем от python-docx

        word = win32com.client.Dispatch("Word.Application") # открываем word
        word.Visible = False # делаем его невидимым

        docx = word.Documents.Open(output_path) # теперь открываем файл через win32com

        style = docx.Styles("Заголовок 1") # Вместо того чтобы создавать новый стиль мы берём дефолтный и изменяем его под наши цели
        style.Font.Name = "Times New Roman"
        style.Font.Size = 14
        style.Font.Color = 0
        style.ParagraphFormat.Alignment = 3
        style.ParagraphFormat.LineSpacingRule = 1 # Интервал 1.5
        cm_to_pt = lambda cm: cm * 72 / 2.54 # Перевод сантиметров в пт
        style.ParagraphFormat.FirstLineIndent = cm_to_pt(1.25) # Красная строка у заголовка
        style.ParagraphFormat.SpaceBefore = 24
        style.ParagraphFormat.SpaceAfter = 24
        style.NoSpaceBetweenParagraphsOfSameStyle = True # Убрать интервал между заголовками

        counters = [0] * 9 # массив для подсчёта заголовков

        for paragraph in docx.Paragraphs: # Проходимся по всем абзацам
            rng = paragraph.Range

            text = rng.Text.rstrip('\r') 

            if text.startswith("~"): # Находим текст начинающийся с ~
                level = min(len(text) - len(text.lstrip("~")), 9) # Считаем количество ~
                cleanText = text.lstrip("~").strip() # убираем ~

                counters[level - 1] += 1 # Записываем этот пункт чтобы он больше не повторялся
                for j in range(level, 9):
                    counters[j] = 0

                number = ".".join(str(counters[k]) for k in range(level) if counters[k] != 0) # Делаем из массива номер
                full_text = f"{number} {cleanText}" # Соединяем номер с текстом

                rng.Text = full_text + "\r" # Заменяем текст на текст с номером
                rng.Style = "Заголовок 1" # Применяем заголовок

                if level == 1 and counters[0] != 1:
                    rng.InsertBefore("\f") #Если это пункт с уровнем 1 то он должен начинаться с нового листа

        if app_data and not use_pdf and not use_check:
            end_rng = docx.Content
            end_rng.Collapse(0)  # Курсор в конец документа

            russian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е']

            for idx, (app_type, images) in enumerate(app_data):
                letter = russian_letters[idx]

                end_rng.InsertBreak(2)  # Новый лист

                para = end_rng.Paragraphs.Add()
                para.Range.Text = f"Приложение {letter}"
                para.Range.Style = "Заголовок 1"
                para.Range.ParagraphFormat.Alignment = 1
                para.Range.ParagraphFormat.SpaceAfter = 0
                para.Range.ParagraphFormat.FirstLineIndent = 0

                para.Range.InsertParagraphAfter()

                end_rng.SetRange(para.Range.End + 1, para.Range.End + 1)

                para2 = end_rng.Paragraphs.Add()
                para2.Range.Text = "(обязательно)"
                para2.Range.Font.Name = "Times New Roman"
                para2.Range.Font.Size = 14
                para2.Range.ParagraphFormat.Alignment = 1

                para2.Range.InsertParagraphAfter()
                end_rng.SetRange(para2.Range.End + 1, para2.Range.End + 1)

                for img_idx, img_path in enumerate(images):
                    # Вставляем абзац для картинки
                    p_img = docx.Paragraphs.Add(end_rng)
                    img_range = p_img.Range
                    img_shape = img_range.InlineShapes.AddPicture(FileName=img_path, LinkToFile=False, SaveWithDocument=True)
                    img_shape.Width = 360  # ширина изображения
                    img_shape.Height = 600 # высота изображения
                    img_shape.LockAspectRatio = True
                    p_img.Alignment = 1  # По центру

                    end_rng.SetRange(p_img.Range.End, p_img.Range.End)

                    # Вставляем подпись сразу после картинки
                    caption_text = f"{app_type} {letter}.{img_idx + 1}"
                    end_rng.InsertAfter(caption_text)

                    # Настраиваем стиль подписи
                    caption_range = end_rng.Duplicate
                    caption_range.SetRange(end_rng.Start, end_rng.Start + len(caption_text))
                    caption_range.Font.Name = "Times New Roman"
                    caption_range.Font.Size = 14
                    caption_range.ParagraphFormat.Alignment = 1  # По центру

                    # После подписи вставляем перенос строки
                    end_rng.SetRange(caption_range.End, caption_range.End)
                    end_rng.InsertParagraphAfter()

                    # И подвигаем курсор для следующего элемента
                    end_rng.SetRange(end_rng.End, end_rng.End)

        if extra_fields:
            lastTitlePara = docx.Paragraphs(titleParagraphsCount - 1)

            rng = lastTitlePara.Range
            rng.Collapse(0) # Ставим курсор в конец титульного листа
            rng.InsertBreak(7) # Перенос на новую страницу

            content = rng.Paragraphs.Add() 
            content.Range.Text = "Содержание"
            content.Range.Font.Name = "Times New Roman"
            content.Range.Font.Size = 14
            content.Range.ParagraphFormat.Alignment = 1  # По центру

            contentRng = content.Range
            contentRng.Collapse(0)  # Ставим курсор после "Содержание"

            # Вставляем Оглавление
            docx.TablesOfContents.Add(
                contentRng,
                UseHeadingStyles=True,
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
                UseFields=False,
                RightAlignPageNumbers=True,
                IncludePageNumbers=True
            )

            toc = docx.TablesOfContents(1)
            toc.Update() # Обновляем оглавление

            # Применяем нужный шрифт и отступ к каждому пункту в оглавлении
            for para in toc.Range.Paragraphs:
                p_rng = para.Range
                p_rng.Font.Name = "Times New Roman"
                p_rng.Font.Size = 14
                p_rng.ParagraphFormat.SpaceAfter = 5
                p_rng.ParagraphFormat.LineSpacingRule = 1

            toc_end_range = docx.TablesOfContents(1).Range
            toc_end_range.Collapse(0)  # Курсор после оглавления
            toc_end_range.InsertBreak(2)  # Перенос на новую страницу

        for section in docx.Sections:
            if extra_fields and section.Index == 1:  # Первая секция (титульная)
                section.PageSetup.DifferentFirstPageHeaderFooter = True # Убираем номера с титульного листа
                first_page_header = section.Headers(1)
                first_page_header.Range.Delete()

            else:  # Остальные секции
                section.PageSetup.DifferentFirstPageHeaderFooter = False
                header = section.Headers(1)
                header.Range.Delete()
                header.Range.Collapse(0)  # Ставим курсор в начало
                field = header.Range.Fields.Add(header.Range, 33)  # Вставляем номер страницы
                header.Range.ParagraphFormat.SpaceAfter = 16 
                field.Result.ParagraphFormat.Alignment = 1  # Центрируем
                field.Result.Font.Name = "Times New Roman"
                field.Result.Font.Size = 14

        docx.Save() # сохранить файл
        docx.Close() # закрыть файл
        word.Quit() # закрыть word

    if use_pdf and not use_check:
        pdf_path = output_path.rsplit('.', 1)[0] + '.pdf' # берём путь с названием файла и добавляем в конец .pdf
        convert(input_path, pdf_path)

    if not use_pdf and use_check:
        doc.save(output_path)

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(output_path)

        for para in doc.Paragraphs:
            for word in para.Range.Words:
                word_text = word.Text.strip()
                if word_text:  # Пропускаем пробелы
                    try:
                        font_name = word.Font.Name
                        font_size = word.Font.Size
                        align = word.ParagraphFormat.Alignment
                        indent = word.ParagraphFormat.FirstLineIndent
                        spacing_rule = word.ParagraphFormat.LineSpacingRule
                    except:
                        continue

                    if align != 3:  # Не выровнен по ширине
                        word.Font.Color = 0x006400
                    if round(indent / 28.35, 2) != 1.25: # Неправильный абзацный отступ
                        word.Font.Color = 0x654321
                    if spacing_rule != WD_LINE_SPACING.ONE_POINT_FIVE:  # не 1.5 интервал
                        word.Font.Color = 0x7B3F99
                    if font_name != "Times New Roman" or int(font_size) != 14: # Не тот шрифт
                        word.Font.Color = 0x000080

        # --- Вставляем лист с расшифровкой ошибок ---
        rng_end = doc.Content
        rng_end.Collapse(0)  # Перемещаемся в самый конец
        rng_end.InsertBreak(2)  # Новый лист (разрыв страницы)

        # Легенда ошибок
        legend_items = [
            ("ошибка шрифта (не Times New Roman, не 14 пт)", 0x000080),
            ("ошибка выравнивания (не по ширине)", 0x006400),
            ("ошибка абзацного отступа (не 1.25 см)", 0x654321),
            ("ошибка межстрочного интервала (не 1.5)", 0x7B3F99),
        ]

        for text, color in legend_items:
            p = rng_end.Paragraphs.Add()
            r = p.Range
            r.Text = text
            r.Font.Name = "Times New Roman"
            r.Font.Size = 14
            p.Alignment = 0  # По левому краю
            r.Font.Color = color

            r.InsertParagraphAfter()
            rng_end.SetRange(r.End + 1, r.End + 1)  # Сдвигаем курсор на новый абзац

        doc.Save()
        doc.Close()
        word.Quit()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('input_path')
    parser.add_argument('output_path')
    parser.add_argument('--pdf', action='store_true', dest='use_pdf')
    parser.add_argument('--check', action='store_true', dest='use_check')
    parser.add_argument('--title', nargs=14)
    parser.add_argument('--appdata', nargs=argparse.REMAINDER)
    
    args = parser.parse_args()

    if args.title:
        extra_fields = []
        for field in args.title:
            extra_fields.append(field.replace('_', ' '))
    else:
        extra_fields = None

    app_data = []
    if args.appdata:
        it = iter(args.appdata)
        apps_count = int(next(it))
        for _ in range(apps_count):
            typ = next(it)
            img_count = int(next(it))
            paths = [next(it) for _ in range(img_count)]
            app_data.append((typ, paths))

    DocxFormat(
        args.input_path,
        args.output_path,
        extra_fields=extra_fields,
        use_pdf=args.use_pdf,
        use_check=args.use_check,
        app_data=app_data
    )

if __name__ == '__main__':
    main()